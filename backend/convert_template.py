"""Prepare the official MPU module-outline files for ``docxtpl``.

The English and Chinese sources are DOCX files. The official Portuguese source
is a legacy DOC file and must first be exported to DOCX by Microsoft Word or
Pages; converting it with ``textutil`` flattens the tables and is not suitable.

Run from the repository root::

    python backend/convert_template.py \
      --source-dir "Module Outline Templates" \
      --pt-docx "Module Outline Templates/module-outline-template_pt_202305.docx"
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


BASE_DIR = Path(__file__).resolve().parent
DEFAULT_SOURCE_DIR = BASE_DIR.parent / "Module Outline Templates"
DEFAULT_OUTPUT_DIR = BASE_DIR / "templates"


def _replace_paragraph_text(paragraph, old: str, new: str) -> bool:
    """Replace text spanning runs while retaining the first run's formatting."""
    if old not in paragraph.text:
        return False
    replacement = paragraph.text.replace(old, new)
    if paragraph.runs:
        paragraph.runs[0].text = replacement
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(replacement)
    return True


def _set_paragraph_text(paragraph, text: str) -> None:
    """Replace visible runs without rebuilding the paragraph or its properties."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _set_cell_text(table, row: int, col: int, text: str) -> None:
    """Replace a value cell while preserving cell and first-run formatting."""
    cell = table.cell(row, col)
    _set_paragraph_text(cell.paragraphs[0], text)
    for extra in cell.paragraphs[1:]:
        _set_paragraph_text(extra, "")


def _fill_metadata_table(table, value_columns: tuple[int, int]) -> None:
    left_value, right_value = value_columns
    mapping = {
        (0, left_value): "{{ academic_year }}",
        (0, right_value): "{{ semester }}",
        (1, left_value): "{{ module_code }}",
        (2, left_value): "{{ module_name }}",
        (3, left_value): "{{ prerequisites }}",
        (4, left_value): "{{ medium_of_instruction }}",
        (5, left_value): "{{ credits }}",
        (5, right_value): "{{ contact_hours }}",
        (6, left_value): "{{ instructor }}",
        (6, right_value): "{{ email }}",
        (7, left_value): "{{ office }}",
        (7, right_value): "{{ office_phone }}",
    }
    for (row, col), placeholder in mapping.items():
        _set_cell_text(table, row, col, placeholder)


def _convert_common(
    src: Path,
    dst: Path,
    header_placeholders: tuple[tuple[str, str], ...],
    required_headers: tuple[str, ...],
    attendance_prefix: str,
    marking_placeholder: str,
    lecturer_placeholders: tuple[str, ...],
    value_columns: tuple[int, int],
) -> None:
    document = Document(src)
    if len(document.tables) != 6:
        raise ValueError(f"{src.name}: expected 6 tables, found {len(document.tables)}")

    replacements_found = {target: False for target in required_headers}
    attendance_found = False
    marking_found = False

    for paragraph in document.paragraphs:
        for source, target in header_placeholders:
            if _replace_paragraph_text(paragraph, source, target):
                replacements_found[target] = True

        stripped = paragraph.text.strip()
        if stripped.startswith(attendance_prefix):
            _set_paragraph_text(paragraph, "{{ attendance_text }}")
            attendance_found = True
        elif marking_placeholder in stripped:
            _set_paragraph_text(paragraph, "__MARKING_RULE_BLOCK__")
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            marking_found = True
        elif stripped in lecturer_placeholders:
            # Retain the official paragraph/style as an editable Word slot, but
            # do not ship instructional placeholder text as lecturer content.
            _set_paragraph_text(paragraph, "")

    missing = [key for key, found in replacements_found.items() if not found]
    if missing or not attendance_found or not marking_found:
        raise ValueError(
            f"{src.name}: required template slots not found "
            f"(headers={missing}, attendance={attendance_found}, marking={marking_found})"
        )

    _fill_metadata_table(document.tables[0], value_columns)
    dst.parent.mkdir(parents=True, exist_ok=True)
    document.save(dst)


def convert_templates(source_dir: Path, output_dir: Path, pt_docx: Path) -> None:
    en_source = source_dir / "module-outline-template_en_202305.docx"
    zh_source = source_dir / "module-outline-template_zh_202305.docx"
    for source in (en_source, zh_source, pt_docx):
        if not source.is_file():
            raise FileNotFoundError(source)

    _convert_common(
        en_source,
        output_dir / "template_en.docx",
        (("[Name of academic unit]", "{{ academic_unit }}"),
         ("[Programme name]", "{{ programme_name }}")),
        ("{{ academic_unit }}", "{{ programme_name }}"),
        "Attendance requirements are governed",
        "[Insert marking scheme]",
        ("[insert text]",),
        (1, 3),
    )
    _convert_common(
        zh_source,
        output_dir / "template_zh.docx",
        (("[學術單位名稱]", "{{ academic_unit }}"),
         ("[課程名稱]", "{{ programme_name }}")),
        ("{{ academic_unit }}", "{{ programme_name }}"),
        "考勤要求按澳門理工大學",
        "[插入評分準則]",
        ("[插入概述]", "[插入書單]", "[插入參考文獻]"),
        (1, 3),
    )

    pt_document = Document(pt_docx)
    pt_columns = len(pt_document.tables[0].columns) if pt_document.tables else 0
    if pt_columns == 5:
        pt_value_columns = (1, 4)
    elif pt_columns == 4:
        pt_value_columns = (1, 3)
    else:
        raise ValueError(
            f"{pt_docx.name}: expected 4 or 5 metadata columns, found {pt_columns}"
        )

    _convert_common(
        pt_docx,
        output_dir / "template_pt.docx",
        (("[nome da unidade académica]", "{{ academic_unit }}"),
         ("[NOME DA UNIDADE ACADÉMICA]", "{{ academic_unit }}"),
         ("[designação do curso]", "{{ programme_name }}"),
         ("[DESIGNAÇÃO DO CURSO]", "{{ programme_name }}")),
        ("{{ academic_unit }}", "{{ programme_name }}"),
        "Os requisitos de assiduidade são cumpridos",
        "[Inserir o critério de classificação]",
        ("[Caracterização]", "[Inserir a bibliografia]", "[Inserir as referências]"),
        pt_value_columns,
    )


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source-dir", type=Path, default=DEFAULT_SOURCE_DIR)
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR)
    parser.add_argument(
        "--pt-docx",
        type=Path,
        default=DEFAULT_SOURCE_DIR / "module-outline-template_pt_202305.docx",
        help="Faithful DOCX export of the official Portuguese .doc file",
    )
    args = parser.parse_args()
    convert_templates(args.source_dir.resolve(), args.output_dir.resolve(), args.pt_docx.resolve())
    print(f"Templates written to {args.output_dir.resolve()}")


if __name__ == "__main__":
    main()
