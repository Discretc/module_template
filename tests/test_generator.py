import io
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "backend"))

import generator  # noqa: E402
from rules import JointRuleConflictError, RULE_PARAGRAPHS, RuleValidationError  # noqa: E402


def sample_class(**overrides):
    data = {
        "class_code": "COMP/9999-001",
        "module_code": "COMP9999",
        "degree_level": "master",
        "faculty_en": "Faculty & Science",
        "faculty_zh": "應用科學學院",
        "faculty_pt": "Faculdade de Ciências Aplicadas",
        "prog_name_en": "Master of Testing",
        "prog_name_zh": "測試碩士學位課程",
        "prog_name_pt": "Mestrado em Testes",
        "module_name_en": "Template <Engineering>",
        "module_name_zh": "範本工程",
        "module_name_pt": "Engenharia de Modelos",
        "prerequisite_en": None,
        "prerequisite_zh": "",
        "prerequisite_pt": None,
        "medium_of_instruction": "English",
        "credits": None,
        "duration": 45,
        "academic_year": "2026/2027",
        "semester": "1",
        "instructor_en": "Ada Lovelace",
        "instructor_zh": None,
        "instructor_pt": "",
        "email": None,
        "room_en": "M505",
        "room_zh": None,
        "room_pt": "",
        "telephone": None,
        "rule_code": 3,
    }
    data.update(overrides)
    return data


def word_xml(doc_bytes):
    with zipfile.ZipFile(io.BytesIO(doc_bytes)) as archive:
        return "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.startswith("word/") and (name.endswith(".xml") or name.endswith(".rels"))
        )


class GeneratorTests(unittest.TestCase):
    def test_templates_retain_official_structure_and_links(self):
        expected_footers = {
            "en": "MPU-LMO-E-v02(2023/06)",
            "zh": "MPU-LMO-C-v02(2023/06)",
            "pt": "MPU-LMO-P-v02(2023/06)",
        }
        for language, path in generator.TEMPLATES.items():
            with self.subTest(language=language):
                document = Document(path)
                self.assertEqual(1, len(document.sections))
                self.assertEqual(6, len(document.tables))
                footer = " ".join(p.text for p in document.sections[0].footer.paragraphs)
                self.assertIn(expected_footers[language], footer)
                package_xml = word_xml(path.read_bytes())
                self.assertIn("assessment_strategy", package_xml)
                self.assertIn("studenthandbook", package_xml)
                self.assertIn("{{ attendance_text }}", package_xml)
                self.assertIn(generator.MARKING_RULE_MARKER, package_xml)

    def test_render_maps_fields_and_preserves_editable_lecturer_slots(self):
        expected = {
            "en": ("Faculty & Science", "Master of Testing", "Master’s", "Nil", "English"),
            "zh": ("應用科學學院", "測試碩士學位課程", "碩士", "無", "中文"),
            "pt": ("Faculdade de Ciências Aplicadas", "Mestrado em Testes", "Mestre", "Não tem", "Português"),
        }
        for language in ("en", "zh", "pt"):
            with self.subTest(language=language):
                rendered = generator._render_one(sample_class(), language)
                document = Document(io.BytesIO(rendered))
                xml = word_xml(rendered)
                faculty, programme, degree, prerequisite, language_label = expected[language]
                self.assertEqual(faculty, document.paragraphs[0].text.strip())
                self.assertEqual(programme, document.paragraphs[1].text.strip())
                self.assertIn(degree, xml)
                self.assertIn("2026/2027", document.tables[0].rows[0].cells[1].text)
                self.assertIn("COMP/9999-001", document.tables[0].rows[1].cells[1].text)
                self.assertEqual(prerequisite, document.tables[0].rows[3].cells[1].text)
                self.assertEqual(language_label, document.tables[0].rows[4].cells[1].text)
                self.assertIn("45", xml)
                for prompt in (
                    "[insert text]", "[插入概述]", "[插入書單]", "[插入參考文獻]",
                    "[Caracterização]", "[Inserir a bibliografia]", "[Inserir as referências]",
                ):
                    self.assertNotIn(prompt, xml)
                for forbidden in ("{{", "}}", "[Doctoral/Master", "[博士/碩士/學士]", "[Doutor / Mestre"):
                    self.assertNotIn(forbidden, xml)
                self.assertNotRegex(xml, r">\s*None\s*<")

    def test_missing_and_invalid_values_render_safely(self):
        empty = sample_class(
            degree_level=None,
            faculty_en=None,
            faculty_zh=None,
            faculty_pt=None,
            prog_name_en=None,
            prog_name_zh=None,
            prog_name_pt=None,
            duration=float("nan"),
            rule_code=1,
        )
        for language in ("en", "zh", "pt"):
            with self.subTest(language=language):
                xml = word_xml(generator._render_one(empty, language))
                self.assertNotRegex(xml, r">\s*None\s*<")
                self.assertNotRegex(xml.lower(), r">\s*nan\s*<")
                self.assertNotIn("{{", xml)
                self.assertNotIn("[Doctoral/Master", xml)
                self.assertNotIn("[博士/碩士/學士]", xml)
                self.assertNotIn("[Doutor / Mestre", xml)

    def test_all_rules_render_as_separate_ordered_paragraphs(self):
        for rule in (1, 2, 3, 4):
            for language in ("en", "zh", "pt"):
                with self.subTest(rule=rule, language=language):
                    expected = RULE_PARAGRAPHS[rule][language]
                    self.assertEqual(
                        expected,
                        generator.get_marking_rule_paragraphs(rule, language),
                    )
                    rendered = generator._render_one(sample_class(rule_code=rule), language)
                    document = Document(io.BytesIO(rendered))
                    paragraphs = [paragraph.text for paragraph in document.paragraphs]
                    for statement in expected:
                        self.assertIn(statement, paragraphs)
                    positions = [paragraphs.index(statement) for statement in expected]
                    self.assertEqual(positions, sorted(positions))
                    self.assertEqual(len(expected), sum(text in expected for text in paragraphs))
                    self.assertNotIn(generator.MARKING_RULE_MARKER, word_xml(rendered))

    def test_rule_one_removes_complete_conditional_paragraph(self):
        document = Document(io.BytesIO(generator._render_one(sample_class(rule_code=1), "en")))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        heading = next(index for index, text in enumerate(paragraphs) if text.casefold() == "marking scheme")
        self.assertEqual("required readings", paragraphs[heading + 1].casefold())

    def test_invalid_rule_codes_are_rejected(self):
        for value in (None, "", "unknown", 0, 5, 2.5):
            with self.subTest(value=value):
                with self.assertRaises(RuleValidationError):
                    generator._render_one(sample_class(rule_code=value), "en")

    def test_conflicting_joint_rules_are_rejected(self):
        joint = sample_class(
            class_codes=["COMP1000-111", "COMP1000-112"],
            rule_codes=[2, 3],
        )
        with self.assertRaisesRegex(
            JointRuleConflictError,
            r"COMP1000-111=Rule 2.*COMP1000-112=Rule 3",
        ):
            generator._render_one(joint, "en")

    def test_batch_zip_names_contents_and_input_immutability(self):
        first = sample_class()
        second = sample_class(class_code="SAFE-002", degree_level="doctoral")
        original = dict(first)
        with tempfile.TemporaryDirectory() as output_dir:
            archive_buffer = generator.generate_batch(
                [first, second],
                academic_year="2027/2028",
                semester="2",
                output_dir=output_dir,
            )
            with zipfile.ZipFile(archive_buffer) as archive:
                names = archive.namelist()
                self.assertEqual(6, len(names))
                self.assertEqual(6, len(set(names)))
                self.assertIn("COMP_9999-001_Module_Outline_EN.docx", names)
                self.assertIn("SAFE-002_Module_Outline_PT.docx", names)
                for name in names:
                    self.assertFalse(name.startswith("/"))
                    self.assertNotIn("..", name)
                    rendered = archive.read(name)
                    self.assertTrue(rendered.startswith(b"PK"))
                    self.assertIn("2027/2028", word_xml(rendered))
            generated = list(Path(output_dir).glob("generated_*/*.docx"))
            self.assertEqual(6, len(generated))
        self.assertEqual(original, first)

    def test_joint_class_codes_are_visible_and_filename_safe(self):
        joint = sample_class(
            class_code="COMP111-111, COMP1121-111, COMP1121-114",
            class_codes=["COMP111-111", "COMP1121-111", "COMP1121-114"],
        )
        rendered = generator._render_one(joint, "en")
        document = Document(io.BytesIO(rendered))
        self.assertEqual(
            "COMP111-111, COMP1121-111, COMP1121-114",
            document.tables[0].rows[1].cells[1].text,
        )
        with tempfile.TemporaryDirectory() as output_dir:
            archive = generator.generate_batch([joint], output_dir=output_dir)
            with zipfile.ZipFile(archive) as zipped:
                self.assertIn(
                    "COMP111-111+COMP1121-111+COMP1121-114_Module_Outline_EN.docx",
                    zipped.namelist(),
                )


class GenerateRouteTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        import app

        cls.app_module = app
        cls.client = app.app.test_client()

    def test_generate_route_returns_batch_zip(self):
        fake_zip = io.BytesIO()
        with zipfile.ZipFile(fake_zip, "w") as archive:
            archive.writestr("sample.docx", b"PK")
        fake_zip.seek(0)

        with patch.object(self.app_module, "get_classes_full", return_value=[sample_class()]) as fetch, \
             patch.object(self.app_module, "generate_batch", return_value=fake_zip) as generate:
            response = self.client.post(
                "/api/generate",
                json={
                    "faculty_id": 1,
                    "programme_id": 2,
                    "class_ids": [7],
                    "academic_year": "2027/2028",
                    "semester": "2",
                },
            )

        self.assertEqual(200, response.status_code)
        self.assertEqual("application/zip", response.mimetype)
        self.assertIn("Module_Outlines.zip", response.headers["Content-Disposition"])
        fetch.assert_called_once_with(class_ids=[7])
        generate.assert_called_once()
        self.assertEqual("2027/2028", generate.call_args.kwargs["academic_year"])
        self.assertEqual("2", generate.call_args.kwargs["semester"])

    def test_generate_route_reports_joint_rule_conflict(self):
        conflict = JointRuleConflictError(
            "Joint class has conflicting Rule values: "
            "COMP1000-111=Rule 2, COMP1000-112=Rule 3"
        )
        with patch.object(self.app_module, "get_classes_full", side_effect=conflict):
            response = self.client.post("/api/generate", json={"class_ids": [7]})

        self.assertEqual(400, response.status_code)
        self.assertIn("COMP1000-111=Rule 2", response.get_json()["error"])
        self.assertIn("COMP1000-112=Rule 3", response.get_json()["error"])


if __name__ == "__main__":
    unittest.main()
